# To add a new cell, type '# %%'
# To add a new markdown cell, type '# %% [markdown]'
# %%
# read text file value
# store each value
# construct to dataframe
# to table
# append that table to text
# text to docx


# %%

from docx import Document
with open(r"./key_attributes.txt", "r") as f:
    data = f.readlines()

months, bankRef, tradeDate, iniValuationDate, effectiveDate, notionalAmount, fixedRate = data[0].split(
    '|', 6)


# %%
print(notionalAmount)


# %%
current = []
t = []
fixed_Rate = []
with open(r"./payment_date.txt", "r") as f:
    data = f.readlines()

print(data[0])

for i in data:
    x, y = i.split('|')
    t.append(x)
    fixed_Rate.append(y.strip('\n'))

print(t)
print(fixed_Rate)


# %%

document = Document()
section = document.sections[0]
header = section.header
header

paragraph = header.paragraphs[0]
paragraph.text = '\tTRADE DETAILS\n\n\tTrade - Term Sheet Template'
paragraph.style = document.styles["Header"]


table = document.add_table(rows=1, cols=2)
table.style = 'Light Grid'
hdr_cells = table.rows[0].cells

first_row = table.add_row().cells
first_row[0].text = "Bank Ref:"
first_row[1].text = bankRef
sec_row = table.add_row().cells
sec_row[0].text = "Trade Date:"
sec_row[1].text = tradeDate
row_3 = table.add_row().cells
row_3[0].text = "Initial Evaluation Date:"
row_3[1].text = iniValuationDate
row_4 = table.add_row().cells
row_4[0].text = "Effective Date:"
row_4[1].text = effectiveDate
row_5 = table.add_row().cells
row_5[0].text = "Physical Settlement Date:"
row_5[1].text = "Scheduled to be the final Fixed Coupon Payment Date, such day being a Clearance System Business Day, subject to occurrence of a Market Disruption Event and/or a Settlement Disruption Event."
row_6 = table.add_row().cells
row_6[0].text = "Maturity Date:"
row_6[1].text = "The Maturity Date shall be the Cash Settlement Date, or if applicable and if later in time to occur due to postponement, the Maturity Date shall be deemed to be the Physical Settlement Date, provided that a Trigger Event has not occurred."
row_7 = table.add_row().cells
row_7[0].text = "Notional Amount:"
row_7[1].text = notionalAmount
row_8 = table.add_row().cells
row_8[0].text = "Fixed Rate"
row_8[1].text = fixedRate

document.add_paragraph()

table2 = document.add_table(rows=1, cols=2)
table2.style = 'Light Grid'
first_row = table2.add_row().cells
first_row[0].text = "(t)"
first_row[1].text = "Fixed Coupon Payment Date(t)"

for (x, y) in zip(t, fixed_Rate):
    row_cells = table2.add_row().cells
    row_cells[0].text = x
    row_cells[1].text = y


document.save('final_template.docx')
